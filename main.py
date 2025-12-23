# main.py
import os
import io
import zipfile
import shutil
import uuid
import json
import random
import schedule
import time
import threading
from datetime import datetime
from typing import List, Dict, Optional
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from groq import Groq
import base64
from io import BytesIO
import requests
from PIL import Image, ImageDraw, ImageFont
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import numpy as np
from pydantic import BaseModel

from utils import (
    save_upload,
    extract_text_from_file,
    generate_worksheets_with_groq,
    create_docx_from_text,
    create_pdf_from_html_optional,
    extract_topic_from_text,
    create_formatted_docx,
    create_formatted_pdf,
    format_grades_display
)

# ---------- Config ----------
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")

# Optional: serve directly from backend folder if templates not used
LANDING_FILE = os.path.join(os.path.dirname(__file__), "landing.html")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
os.makedirs(TEMPLATE_DIR, exist_ok=True)

# ---------- Groq API ----------

from dotenv import load_dotenv
import os

load_dotenv()
API_KEY = os.getenv("GROQ_API_KEY")
client = Groq(api_key=API_KEY)

# ---------- App ----------
app = FastAPI(title="EduAssist (Groq) Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # limit in production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------- In-memory store ----------
conversations = {}
student_records = {}
progress_history = {}

# ---------- Pydantic Models ----------
class StudentRecord(BaseModel):
    student_id: str
    name: str
    class_name: str
    grade: str
    progress: float
    status: str
    attendance: float
    last_updated: str

class ProgressUpdate(BaseModel):
    student_id: str
    class_name: str
    assignment_name: str
    score: float
    max_score: float
    assignment_type: str
    date_completed: str
    notes: Optional[str] = None

class StudentCreate(BaseModel):
    name: str
    class_name: str
    grade: str
    email: Optional[str] = None
    parent_contact: Optional[str] = None

# ---------- Serve frontend pages ----------
@app.get("/", response_class=HTMLResponse)
async def get_home():
    file_path = os.path.join(os.path.dirname(__file__), "landing.html")
    if os.path.exists(file_path):
        return FileResponse(file_path)
    return HTMLResponse("<h1>Landing page not found.</h1>", status_code=404)

@app.get("/lesson-plan", response_class=HTMLResponse)
async def serve_lesson_plan():
    file_path = os.path.join(os.path.dirname(__file__), "lesson_planning.html")
    if os.path.exists(file_path):
        return FileResponse(file_path)
    return HTMLResponse("<h2>Lesson Plan page not found</h2>", status_code=404)

@app.get("/differentiated-materials", response_class=HTMLResponse)
async def serve_differentiated_materials():
    file_path = os.path.join(os.path.dirname(__file__), "Differentiated_materials.html")
    if os.path.exists(file_path):
        return FileResponse(file_path)
    return HTMLResponse("<h2>Differentiated Materials page not found</h2>", status_code=404)

@app.get("/fun-activities", response_class=HTMLResponse)
async def serve_fun_activities():
    file_path = os.path.join(os.path.dirname(__file__), "fun_activities.html")
    if os.path.exists(file_path):
        return FileResponse(file_path)
    return HTMLResponse("<h2>Fun Activities page not found</h2>", status_code=404)

@app.get("/blackboard-diagrams", response_class=HTMLResponse)
async def serve_blackboard_diagrams():
    file_path = os.path.join(os.path.dirname(__file__), "blackboard_diagrams.html")
    if os.path.exists(file_path):
        return FileResponse(file_path)
    return HTMLResponse("<h2>Blackboard Diagrams page not found</h2>", status_code=404)

@app.get("/student-progress", response_class=HTMLResponse)
async def serve_student_progress():
    """
    Serve the student progress tracking page
    """
    file_path = os.path.join(os.path.dirname(__file__), "Student_tracking.html")
    if os.path.exists(file_path):
        return FileResponse(file_path)
    return HTMLResponse("<h2>Student Progress page not found</h2>", status_code=404)

@app.get("/knowledge-base", response_class=HTMLResponse)
async def serve_knowledge_base():
    file_path = os.path.join(os.path.dirname(__file__), "knowledge_base.html")
    if os.path.exists(file_path):
        return FileResponse(file_path)
    return HTMLResponse("<h2>Knowledge Base page not found</h2>", status_code=404)

# ---------- Lesson Plan Generator ----------
@app.post("/generate/")
async def generate(request: Request):
    try:
        data = await request.json()
        prompt = data.get("prompt", "")
        grade = data.get("grade")
        conversation_id = data.get("conversation_id") or str(uuid.uuid4())

        if conversation_id not in conversations:
            conversations[conversation_id] = {
                "id": conversation_id,
                "created_at": datetime.now().isoformat(),
                "messages": [],
                "lesson_plans": []
            }

        user_message = {
            "role": "user",
            "content": prompt,
            "timestamp": datetime.now().isoformat(),
            "grade": grade
        }
        conversations[conversation_id]["messages"].append(user_message)

        is_lesson_request = any(keyword in prompt.lower() for keyword in [
            "lesson plan", "lesson", "teaching plan", "class plan",
            "curriculum", "syllabus", "teaching", "educate", "teach",
            "classroom activity", "learning objectives", "lesson outline"
        ])

        if is_lesson_request:
            if not grade:
                bot_response = {
                    "response": "I'd be happy to create a lesson plan for you! For which grade level would you like me to design this lesson?",
                    "type": "grade_request",
                    "conversation_id": conversation_id,
                    "requires_grade": True
                }
                conversations[conversation_id]["messages"].append({
                    "role": "assistant",
                    "content": bot_response["response"],
                    "timestamp": datetime.now().isoformat(),
                    "type": "grade_request"
                })
                return JSONResponse(content=bot_response)

            topic = prompt.lower()
            for phrase in ["lesson plan", "lesson", "create", "make", "generate", "for grade", "about"]:
                topic = topic.replace(phrase, "")
            topic = topic.strip().capitalize() or "General Topic"

            lesson_prompt = f"""
You are an expert educator. Create a comprehensive, practical, and engaging lesson plan for teachers.

GRADE LEVEL: {grade}
TOPIC: {topic}
DURATION: 45-60 minutes

Format response with the following sections:
🎯 LESSON OVERVIEW
📚 LEARNING OBJECTIVES
🧩 MATERIALS & RESOURCES
⏰ LESSON PROCEDURE
🎨 DIFFERENTIATION STRATEGIES
📝 ASSESSMENT & FEEDBACK
💡 TEACHER TIPS

Keep it:
- Teacher-friendly
- Age-appropriate for grade {grade}
- Clear, engaging, and visually structured with emojis
- Practical and actionable
"""

            resp = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": lesson_prompt}],
                temperature=0.7,
                max_tokens=2000
            )

            lesson_content = getattr(resp.choices[0].message, "content", str(resp))
            lesson_plan = {
                "id": str(uuid.uuid4()),
                "topic": topic,
                "grade": grade,
                "content": lesson_content,
                "created_at": datetime.now().isoformat(),
                "conversation_id": conversation_id
            }
            conversations[conversation_id]["lesson_plans"].append(lesson_plan)

            bot_response = {
                "response": lesson_content,
                "type": "lesson_plan",
                "grade": grade,
                "topic": topic,
                "conversation_id": conversation_id,
                "lesson_id": lesson_plan["id"]
            }
        else:
            conversation_history = conversations[conversation_id]["messages"][-6:]
            messages_for_api = [
                {"role": "system", "content": """
You are a helpful AI teaching assistant. 
Support teachers by providing actionable classroom ideas, 
teaching methods, and age-appropriate advice. 
If they ask for a lesson plan, ask their grade level first.
Be practical, encouraging, and focus on real classroom applications.
"""}
            ]
            for msg in conversation_history:
                messages_for_api.append({"role": msg["role"], "content": msg["content"]})

            messages_for_api.append({"role": "user", "content": prompt})

            resp = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=messages_for_api,
                temperature=0.7,
                max_tokens=1000
            )

            chat_content = getattr(resp.choices[0].message, "content", str(resp))
            bot_response = {
                "response": chat_content,
                "type": "chat",
                "conversation_id": conversation_id
            }

        conversations[conversation_id]["messages"].append({
            "role": "assistant",
            "content": bot_response["response"],
            "timestamp": datetime.now().isoformat(),
            "type": bot_response["type"]
        })

        print(f"\n✅ [{bot_response['type'].upper()}] Response generated for conversation {conversation_id} at {datetime.now()}")
        return JSONResponse(content=bot_response)

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating response: {str(e)}")

# ---------- PDF Summarizer Endpoint ----------
@app.post("/summarize-pdf/")
async def summarize_pdf(
    file: UploadFile = File(...),
    summary_length: str = Form("medium"),
    page_start: int = Form(1),
    page_end: int = Form(None)
):
    """
    Dedicated endpoint for PDF summarization with page range support
    """
    try:
        # Save uploaded file
        saved_path = await save_upload(file, UPLOAD_FOLDER)
        if not saved_path:
            raise HTTPException(status_code=400, detail="Failed to save uploaded file")
        
        # Extract text from file with page range
        source_text = extract_text_from_file(saved_path, page_start=page_start, page_end=page_end)
        if not source_text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the specified pages")
        
        # Determine summary length parameters
        length_params = {
            "short": {"max_tokens": 500, "temperature": 0.7},
            "medium": {"max_tokens": 1000, "temperature": 0.7},
            "detailed": {"max_tokens": 2000, "temperature": 0.7}
        }
        
        params = length_params.get(summary_length, length_params["medium"])
        
        # Create summarization prompt with page range info
        page_range_info = f"Pages {page_start}" + (f"-{page_end}" if page_end else "+")
        summarization_prompt = f"""
Please provide a comprehensive and well-structured summary of the following document content from {page_range_info}.
Focus on the main ideas, key concepts, and important details.

DOCUMENT CONTENT:
{source_text[:12000]}  # Limit context to avoid token limits

Please structure your summary with:
- Main topic and purpose
- Key points and concepts
- Important findings or conclusions
- Relevance and applications

Make the summary {summary_length} in length, focusing on clarity and coherence.
Avoid repeating the same information multiple times.
Provide a natural flow from introduction to conclusion.
"""
        
        # Generate summary using Groq
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system", 
                    "content": "You are an expert at summarizing educational content. Create clear, concise, and well-structured summaries that capture the essence of the document while maintaining accuracy and readability. Avoid repetition and ensure the summary flows naturally from introduction to conclusion."
                },
                {
                    "role": "user",
                    "content": summarization_prompt
                }
            ],
            temperature=params["temperature"],
            max_tokens=params["max_tokens"],
            top_p=0.9
        )
        
        summary_content = response.choices[0].message.content
        
        # Create output file
        job_id = str(uuid.uuid4())
        output_dir = os.path.join(OUTPUT_FOLDER, job_id)
        os.makedirs(output_dir, exist_ok=True)
        
        # Save summary as PDF only (no ZIP)
        pdf_filename = f"summary_{summary_length}_{job_id[:8]}.pdf"
        pdf_path = os.path.join(OUTPUT_FOLDER, pdf_filename)
        
        create_pdf_from_html_optional(f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; margin: 40px; }}
                h1 {{ color: #333; border-bottom: 2px solid #6C63FF; padding-bottom: 10px; }}
                .metadata {{ background: #f5f5f5; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
                .content {{ margin-top: 20px; }}
                .section {{ margin-bottom: 20px; }}
            </style>
        </head>
        <body>
            <h1>Document Summary</h1>
            <div class="metadata">
                <strong>Original File:</strong> {file.filename}<br>
                <strong>Summary Length:</strong> {summary_length}<br>
                <strong>Pages:</strong> {page_range_info}<br>
                <strong>Generated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M')}
            </div>
            <div class="content">
                {summary_content.replace('\n', '<br>')}
            </div>
        </body>
        </html>
        """, pdf_path)
        
        return JSONResponse(content={
            "summary": summary_content,
            "download_url": f"/download/{pdf_filename}",
            "filename": pdf_filename,
            "page_range": page_range_info
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Summarization failed: {str(e)}")

# ---------- IMPROVED Worksheet Generator ----------
@app.post("/generate-worksheets/")
async def generate_worksheets(
    file: UploadFile = File(...),
    subject: str = Form(None),
    grades: str = Form(...),
    difficulty: str = Form("moderate"),
    question_types: str = Form("mcq,theory"),  # New parameter for question types
    output_format: str = Form("pdf"),
    page_start: int = Form(1),
    page_end: int = Form(None)
):
    """
    Improved worksheet generator with proper question type formatting
    """
    try:
        # Save uploaded file
        saved_path = await save_upload(file, UPLOAD_FOLDER)
        if not saved_path:
            raise HTTPException(status_code=400, detail="Failed to save uploaded file")
        
        # Extract text from file with page range
        source_text = extract_text_from_file(saved_path, page_start=page_start, page_end=page_end)
        if not source_text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the specified pages")
        
        # Extract topic from source text if subject not provided
        extracted_topic = extract_topic_from_text(source_text)
        actual_subject = subject or extracted_topic
        
        # Parse grades and format nicely
        grade_list = [g.strip() for g in grades.split(",") if g.strip()]
        if not grade_list:
            raise HTTPException(status_code=400, detail="No valid grades provided")
        
        # Parse question types
        question_type_list = [q.strip().lower() for q in question_types.split(",") if q.strip()]
        if not question_type_list:
            question_type_list = ["mcq", "theory"]
        
        grades_display = format_grades_display(grade_list)
        
        # Generate worksheets
        worksheets = generate_worksheets_with_groq(
            source_text=source_text,
            subject=actual_subject,
            grades=grade_list,
            difficulty=difficulty,
            question_types=question_type_list,  # Pass question types
            groq_client=client
        )
        
        # Create download links for each grade
        download_links = {}
        page_range_info = f"Pages {page_start}" + (f"-{page_end}" if page_end else "+")
        
        for grade, content in worksheets.items():
            # Create safe filename
            safe_topic = "".join(c for c in actual_subject if c.isalnum() or c in (' ', '-', '_')).rstrip()
            safe_topic = safe_topic.replace(' ', '_')[:30]
            
            # Include question types in filename
            qtypes_str = "_".join(question_type_list)
            filename_base = f"{safe_topic}_Grade_{grade}_{difficulty}_{qtypes_str}"
            
            pdf_filename = f"{filename_base}.pdf"
            pdf_path = os.path.join(OUTPUT_FOLDER, pdf_filename)
            
            docx_filename = f"{filename_base}.docx"
            docx_path = os.path.join(OUTPUT_FOLDER, docx_filename)
            
            # Create formatted files
            create_formatted_pdf(content, actual_subject, grade, difficulty, pdf_path, question_type_list)
            create_formatted_docx(content, actual_subject, grade, difficulty, docx_path, question_type_list)
            
            download_links[grade] = {
                'pdf': f"/download/{pdf_filename}",
                'docx': f"/download/{docx_filename}",
                'preview': content[:300] + "..." if len(content) > 300 else content
            }
        
        # Clean up uploaded file
        try:
            os.remove(saved_path)
        except:
            pass
        
        return JSONResponse(content={
            "status": "success",
            "topic": actual_subject,
            "grades_display": grades_display,
            "page_range": page_range_info,
            "question_types": question_type_list,
            "worksheets": download_links,
            "message": f"Generated {len(grade_list)} worksheet(s) for {grades_display} with {', '.join(question_type_list)} questions"
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Worksheet generation failed: {str(e)}")

# ---------- Topic-based Worksheet Generator ----------
@app.post("/generate-from-topics/")
async def generate_from_topics(
    main_topic: str = Form(...),
    subject: str = Form(...),
    grades: str = Form(...),
    difficulty: str = Form("moderate"),
    question_types: str = Form("mcq,theory"),
    subtopics: str = Form(""),  # Make this optional with default
    output_format: str = Form("pdf")
):
    """
    Generate worksheets directly from topics without file upload
    """
    try:
        print(f"Received topic generation request: {main_topic}, {subject}, {grades}")  # Debug log
        
        if not main_topic.strip():
            raise HTTPException(status_code=400, detail="Main topic is required")
        
        if not subject.strip():
            raise HTTPException(status_code=400, detail="Subject is required")
        
        grade_list = [g.strip() for g in grades.split(",") if g.strip()]
        if not grade_list:
            raise HTTPException(status_code=400, detail="At least one grade level is required")
        
        # Parse question types
        question_type_list = [q.strip().lower() for q in question_types.split(",") if q.strip()]
        if not question_type_list:
            question_type_list = ["mcq", "theory"]
        
        subtopic_list = [s.strip() for s in subtopics.split(",") if s.strip()] if subtopics else []
        
        print(f"Processing: {len(grade_list)} grades, {len(question_type_list)} question types")  # Debug log
        
        # Create source text from topics
        source_text = f"""
MAIN TOPIC: {main_topic}
SUBTOPICS: {', '.join(subtopic_list) if subtopic_list else 'All relevant aspects'}
SUBJECT: {subject}
GRADE LEVELS: {', '.join(grade_list)}
DIFFICULTY: {difficulty}
QUESTION TYPES: {', '.join(question_type_list)}

Please create comprehensive educational worksheets focusing on {main_topic}.
"""
        
        if subtopic_list:
            source_text += f"\nSPECIFICALLY COVER THESE SUBTOPICS: {', '.join(subtopic_list)}"
        
        # Generate worksheets
        worksheets = generate_worksheets_with_groq(
            source_text=source_text,
            subject=subject,
            grades=grade_list,
            difficulty=difficulty,
            question_types=question_type_list,
            groq_client=client
        )
        
        # Create download links
        download_links = {}
        safe_topic = "".join(c for c in main_topic if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_topic = safe_topic.replace(' ', '_')[:30]
        
        for grade, content in worksheets.items():
            # Include question types in filename
            qtypes_str = "_".join(question_type_list)
            filename_base = f"{safe_topic}_Grade_{grade}_{difficulty}_{qtypes_str}"
            
            pdf_filename = f"{filename_base}.pdf"
            pdf_path = os.path.join(OUTPUT_FOLDER, pdf_filename)
            
            docx_filename = f"{filename_base}.docx"
            docx_path = os.path.join(OUTPUT_FOLDER, docx_filename)
            
            create_formatted_pdf(content, subject, grade, difficulty, pdf_path, question_type_list)
            create_formatted_docx(content, subject, grade, difficulty, docx_path, question_type_list)
            
            download_links[grade] = {
                'pdf': f"/download/{pdf_filename}",
                'docx': f"/download/{docx_filename}",
                'preview': content[:300] + "..." if len(content) > 300 else content
            }
        
        return JSONResponse(content={
            "status": "success", 
            "topic": main_topic,
            "subject": subject,
            "grades_display": format_grades_display(grade_list),
            "question_types": question_type_list,
            "worksheets": download_links,
            "subtopics": subtopic_list,
            "message": f"Successfully generated worksheets for {main_topic}"
        })
        
    except Exception as e:
        import traceback
        print(f"Error in generate-from-topics: {str(e)}")  # Debug log
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Topic-based generation failed: {str(e)}")

# ---------- Subtopics Generator ----------
@app.post("/generate-subtopics/")
async def generate_subtopics(request: Request):
    """
    Generate relevant subtopics for a main topic
    """
    try:
        data = await request.json()
        main_topic = data.get("main_topic", "").strip()
        
        if not main_topic:
            raise HTTPException(status_code=400, detail="Main topic is required")
        
        # Generate subtopics using Groq
        subtopic_prompt = f"""
Given the main topic "{main_topic}", generate 5-7 relevant subtopics that would be appropriate for educational worksheets.
Return only a comma-separated list of subtopics without any additional text, numbering, or explanations.

Example format: Algebra, Geometry, Calculus, Statistics, Trigonometry

Focus on subtopics that are:
- Educational and teachable
- Appropriate for worksheet creation
- Cover different aspects of the main topic
- Suitable for various grade levels
"""
        
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": "You are an educational expert. Generate relevant, educational subtopics suitable for worksheet creation. Return only a comma-separated list without any additional text."
                },
                {
                    "role": "user",
                    "content": subtopic_prompt
                }
            ],
            temperature=0.7,
            max_tokens=200
        )
        
        subtopics_text = response.choices[0].message.content.strip()
        # Clean up the response
        subtopics_text = subtopics_text.replace('.', '').replace('\n', ',')
        subtopics_list = [s.strip() for s in subtopics_text.split(',') if s.strip()]
        
        # Remove duplicates and limit to 7
        unique_subtopics = []
        for subtopic in subtopics_list:
            if subtopic not in unique_subtopics and len(unique_subtopics) < 7:
                unique_subtopics.append(subtopic)
        
        return JSONResponse(content={
            "main_topic": main_topic,
            "subtopics": unique_subtopics
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate subtopics: {str(e)}")

# ---------- Fun Learning Activities Generator ----------
@app.post("/generate-activity/")
async def generate_activity(
    subject: str = Form(...),
    topic: str = Form(...),
    grade_level: str = Form(...),
    activity_type: str = Form(None),
    duration: str = Form(None),
    resources: str = Form(""),
    learning_objective: str = Form("")
):
    """
    Generate fun learning activities and games based on curriculum and resources
    """
    try:
        if not subject or not topic or not grade_level:
            raise HTTPException(status_code=400, detail="Subject, topic, and grade level are required")
        
        # Create activity generation prompt
        activity_prompt = f"""
Create a fun, engaging learning activity or game for teachers to use in the classroom.

SUBJECT: {subject}
TOPIC: {topic}
GRADE LEVEL: {grade_level}
LEARNING OBJECTIVE: {learning_objective or "Make learning engaging and memorable"}
AVAILABLE RESOURCES: {resources or "Standard classroom materials"}
DURATION: {duration or "Flexible"}
ACTIVITY TYPE: {activity_type or "Any engaging format"}

Please provide a comprehensive activity plan with:

1. ACTIVITY TITLE: Creative and engaging name
2. ACTIVITY TYPE: Game, group activity, individual task, etc.
3. DURATION: Estimated time required
4. MATERIALS NEEDED: List of required resources
5. STEP-BY-STEP INSTRUCTIONS: Clear, actionable steps for teachers
6. LEARNING OUTCOMES: What students will learn/practice
7. DIFFERENTIATION STRATEGIES: How to adapt for different learners
8. ASSESSMENT IDEAS: How to measure learning
9. VARIATIONS: Alternative ways to run the activity

Make the activity:
- Highly engaging and interactive
- Age-appropriate for the grade level
- Connected to the learning objective
- Practical with available resources
- Easy to implement for teachers
- Fun and memorable for students

Focus on active learning, student participation, and making the content come alive.
"""

        # Generate activity using Groq
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": """You are an expert educational game designer and activity creator. 
                    Create fun, engaging, and educational activities that make learning enjoyable. 
                    Focus on interactive, hands-on experiences that promote active learning.
                    Provide practical, ready-to-use activities that teachers can implement easily."""
                },
                {
                    "role": "user",
                    "content": activity_prompt
                }
            ],
            temperature=0.8,
            max_tokens=3000,
            top_p=0.9
        )
        
        activity_content = response.choices[0].message.content
        
        # Create a formatted response
        return JSONResponse(content={
            "status": "success",
            "activity": activity_content,
            "subject": subject,
            "topic": topic,
            "grade_level": grade_level,
            "activity_type": activity_type or "mixed",
            "message": f"Created fun activity for {topic} in {subject}"
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Activity generation failed: {str(e)}")

# ---------- Student Progress Tracking Endpoints ----------

@app.get("/api/students/")
async def get_all_students(
    class_filter: str = None,
    status_filter: str = None,
    grade_filter: str = None
):
    """
    Get all students with optional filtering
    """
    try:
        filtered_students = {}
        
        for student_id, student in student_records.items():
            # Apply filters
            if class_filter and class_filter != "all" and student.get('class_name') != class_filter:
                continue
            if status_filter and status_filter != "all" and student.get('status') != status_filter:
                continue
            if grade_filter and grade_filter != "all" and student.get('grade') != grade_filter:
                continue
                
            filtered_students[student_id] = student
        
        return JSONResponse(content={
            "students": filtered_students,
            "total_count": len(filtered_students),
            "filters_applied": {
                "class": class_filter,
                "status": status_filter,
                "grade": grade_filter
            }
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching students: {str(e)}")

@app.post("/api/students/")
async def create_student(student_data: StudentCreate):
    """
    Create a new student record
    """
    try:
        student_id = str(uuid.uuid4())
        
        # Initialize student record
        student_record = {
            "student_id": student_id,
            "name": student_data.name,
            "class_name": student_data.class_name,
            "grade": student_data.grade,
            "email": student_data.email,
            "parent_contact": student_data.parent_contact,
            "progress": 0.0,
            "attendance": 100.0,  # Start with perfect attendance
            "status": "average",
            "last_updated": datetime.now().isoformat(),
            "created_at": datetime.now().isoformat()
        }
        
        student_records[student_id] = student_record
        
        # Initialize progress history
        progress_history[student_id] = []
        
        return JSONResponse(content={
            "status": "success",
            "student_id": student_id,
            "message": f"Student {student_data.name} created successfully"
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error creating student: {str(e)}")

@app.get("/api/students/{student_id}")
async def get_student_details(student_id: str):
    """
    Get detailed information for a specific student
    """
    try:
        if student_id not in student_records:
            raise HTTPException(status_code=404, detail="Student not found")
        
        student = student_records[student_id]
        history = progress_history.get(student_id, [])
        
        # Calculate statistics
        recent_scores = [entry['percentage'] for entry in history[-10:]]  # Last 10 assignments
        average_score = sum(recent_scores) / len(recent_scores) if recent_scores else 0
        
        return JSONResponse(content={
            "student": student,
            "progress_history": history,
            "statistics": {
                "average_score": round(average_score, 2),
                "total_assignments": len(history),
                "recent_trend": calculate_trend(history),
                "attendance_rate": student.get('attendance', 100)
            }
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error fetching student details: {str(e)}")

@app.post("/api/progress/")
async def update_student_progress(progress_data: ProgressUpdate):
    """
    Update student progress with new assignment/test results
    """
    try:
        student_id = progress_data.student_id
        
        if student_id not in student_records:
            raise HTTPException(status_code=404, detail="Student not found")
        
        # Calculate percentage score
        percentage = (progress_data.score / progress_data.max_score) * 100
        
        # Create progress entry
        progress_entry = {
            "entry_id": str(uuid.uuid4()),
            "student_id": student_id,
            "assignment_name": progress_data.assignment_name,
            "assignment_type": progress_data.assignment_type,
            "score": progress_data.score,
            "max_score": progress_data.max_score,
            "percentage": round(percentage, 2),
            "date_completed": progress_data.date_completed,
            "notes": progress_data.notes,
            "recorded_at": datetime.now().isoformat()
        }
        
        # Add to history
        if student_id not in progress_history:
            progress_history[student_id] = []
        progress_history[student_id].append(progress_entry)
        
        # Update student's overall progress and status
        await update_student_overall_progress(student_id)
        
        return JSONResponse(content={
            "status": "success",
            "message": "Progress updated successfully",
            "progress_entry": progress_entry
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating progress: {str(e)}")

@app.get("/api/progress/overview")
async def get_progress_overview(
    class_name: str = None,
    time_range: str = "current_semester"
):
    """
    Get overview statistics and charts data
    """
    try:
        # Filter students by class
        filtered_students = {}
        for student_id, student in student_records.items():
            if class_name and class_name != "all" and student.get('class_name') != class_name:
                continue
            filtered_students[student_id] = student
        
        # Calculate statistics
        total_students = len(filtered_students)
        if total_students == 0:
            return JSONResponse(content={
                "total_students": 0,
                "average_progress": 0,
                "attendance_rate": 0,
                "grade_distribution": {},
                "performance_trends": []
            })
        
        # Calculate averages
        total_progress = sum(student.get('progress', 0) for student in filtered_students.values())
        total_attendance = sum(student.get('attendance', 0) for student in filtered_students.values())
        
        average_progress = total_progress / total_students
        average_attendance = total_attendance / total_students
        
        # Calculate grade distribution
        grade_distribution = {
            "excellent": 0,  # 90-100%
            "good": 0,       # 80-89%
            "average": 0,    # 70-79%
            "poor": 0        # <70%
        }
        
        for student in filtered_students.values():
            progress = student.get('progress', 0)
            if progress >= 90:
                grade_distribution["excellent"] += 1
            elif progress >= 80:
                grade_distribution["good"] += 1
            elif progress >= 70:
                grade_distribution["average"] += 1
            else:
                grade_distribution["poor"] += 1
        
        # Generate performance trends (last 8 weeks)
        performance_trends = generate_performance_trends(filtered_students, time_range)
        
        return JSONResponse(content={
            "total_students": total_students,
            "average_progress": round(average_progress, 2),
            "attendance_rate": round(average_attendance, 2),
            "grade_distribution": grade_distribution,
            "performance_trends": performance_trends
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating overview: {str(e)}")

@app.post("/api/progress/bulk-update")
async def bulk_update_progress(updates: List[ProgressUpdate]):
    """
    Update progress for multiple students at once
    """
    try:
        results = []
        for update in updates:
            try:
                # Reuse individual update logic
                student_id = update.student_id
                
                if student_id not in student_records:
                    results.append({
                        "student_id": student_id,
                        "status": "error",
                        "message": "Student not found"
                    })
                    continue
                
                percentage = (update.score / update.max_score) * 100
                
                progress_entry = {
                    "entry_id": str(uuid.uuid4()),
                    "student_id": student_id,
                    "assignment_name": update.assignment_name,
                    "assignment_type": update.assignment_type,
                    "score": update.score,
                    "max_score": update.max_score,
                    "percentage": round(percentage, 2),
                    "date_completed": update.date_completed,
                    "notes": update.notes,
                    "recorded_at": datetime.now().isoformat()
                }
                
                if student_id not in progress_history:
                    progress_history[student_id] = []
                progress_history[student_id].append(progress_entry)
                
                await update_student_overall_progress(student_id)
                
                results.append({
                    "student_id": student_id,
                    "status": "success",
                    "message": "Progress updated"
                })
                
            except Exception as e:
                results.append({
                    "student_id": update.student_id,
                    "status": "error",
                    "message": str(e)
                })
        
        return JSONResponse(content={
            "results": results,
            "total_processed": len(updates),
            "successful": len([r for r in results if r['status'] == 'success'])
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error in bulk update: {str(e)}")

@app.post("/api/reports/generate")
async def generate_progress_report(
    report_type: str = Form("class_summary"),
    class_name: str = Form(None),
    student_ids: str = Form(None),
    include_charts: bool = Form(True),
    format: str = Form("pdf")
):
    """
    Generate comprehensive progress reports
    """
    try:
        # Parse student IDs if provided
        target_student_ids = []
        if student_ids:
            target_student_ids = [sid.strip() for sid in student_ids.split(",") if sid.strip()]
        
        # Get data based on report type
        if report_type == "class_summary":
            report_data = await generate_class_summary_report(class_name, include_charts)
        elif report_type == "individual_reports":
            report_data = await generate_individual_reports(target_student_ids, include_charts)
        elif report_type == "comparison_report":
            report_data = await generate_comparison_report(target_student_ids, include_charts)
        else:
            raise HTTPException(status_code=400, detail="Invalid report type")
        
        # Generate file
        job_id = str(uuid.uuid4())
        filename = f"progress_report_{report_type}_{job_id[:8]}.{format}"
        filepath = os.path.join(OUTPUT_FOLDER, filename)
        
        if format == "pdf":
            create_progress_report_pdf(report_data, filepath)
        else:
            create_progress_report_docx(report_data, filepath)
        
        return JSONResponse(content={
            "status": "success",
            "download_url": f"/download/{filename}",
            "filename": filename,
            "report_type": report_type,
            "generated_at": datetime.now().isoformat()
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating report: {str(e)}")

@app.get("/api/analytics/performance-trends")
async def get_performance_trends(
    class_name: str = None,
    weeks: int = 8
):
    """
    Get performance trends data for charts
    """
    try:
        trends_data = generate_performance_trends(student_records, f"last_{weeks}_weeks")
        
        return JSONResponse(content={
            "trends": trends_data,
            "time_period": f"Last {weeks} weeks",
            "class": class_name or "All Classes"
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating trends: {str(e)}")

# ---------- Helper Functions for Student Progress ----------

async def update_student_overall_progress(student_id: str):
    """
    Update student's overall progress based on recent performance
    """
    try:
        if student_id not in progress_history or not progress_history[student_id]:
            return
        
        # Get recent assignments (last 10 or all if less)
        recent_entries = progress_history[student_id][-10:]
        
        # Calculate weighted average (recent assignments matter more)
        total_weight = 0
        weighted_sum = 0
        
        for i, entry in enumerate(recent_entries):
            weight = i + 1  # More recent = higher weight
            total_weight += weight
            weighted_sum += entry['percentage'] * weight
        
        new_progress = weighted_sum / total_weight if total_weight > 0 else 0
        
        # Update student record
        student_records[student_id]['progress'] = round(new_progress, 2)
        student_records[student_id]['last_updated'] = datetime.now().isoformat()
        
        # Update status based on progress
        if new_progress >= 90:
            status = "excellent"
        elif new_progress >= 80:
            status = "good"
        elif new_progress >= 70:
            status = "average"
        else:
            status = "poor"
        
        student_records[student_id]['status'] = status
        
    except Exception as e:
        print(f"Error updating overall progress for {student_id}: {e}")

def calculate_trend(history: List[Dict]) -> str:
    """
    Calculate performance trend (improving, declining, stable)
    """
    if len(history) < 2:
        return "stable"
    
    recent = history[-5:]  # Last 5 entries
    if len(recent) < 2:
        return "stable"
    
    first_avg = sum(entry['percentage'] for entry in recent[:2]) / 2
    last_avg = sum(entry['percentage'] for entry in recent[-2:]) / 2
    
    if last_avg - first_avg > 5:
        return "improving"
    elif first_avg - last_avg > 5:
        return "declining"
    else:
        return "stable"

def generate_performance_trends(students: Dict, time_range: str) -> List[Dict]:
    """
    Generate performance trends data for charts
    """
    trends = []
    weeks = 8
    
    for week in range(weeks):
        week_data = {
            "week": f"Week {week + 1}",
            "average_score": 70 + (week * 2) + (random.random() * 10 - 5),  # Simulated improvement
            "top_performer": 95 + (random.random() * 5 - 2.5),
            "class_average": 75 + (week * 1.5) + (random.random() * 8 - 4),
            "assignments_completed": len(students) * (0.8 + (week * 0.05))  # Simulated increase
        }
        trends.append(week_data)
    
    return trends

async def generate_class_summary_report(class_name: str, include_charts: bool) -> Dict:
    """
    Generate class summary report data
    """
    filtered_students = {
        sid: student for sid, student in student_records.items()
        if not class_name or student.get('class_name') == class_name
    }
    
    overview_data = await get_progress_overview(class_name, "current_semester")
    overview_dict = json.loads(overview_data.body) if hasattr(overview_data, 'body') else overview_data
    
    return {
        "report_type": "class_summary",
        "class_name": class_name or "All Classes",
        "generated_at": datetime.now().isoformat(),
        "summary": overview_dict,
        "students": filtered_students,
        "include_charts": include_charts
    }

async def generate_individual_reports(student_ids: List[str], include_charts: bool) -> Dict:
    """
    Generate individual student reports
    """
    individual_reports = {}
    
    for student_id in student_ids:
        if student_id in student_records:
            student_data = await get_student_details(student_id)
            individual_reports[student_id] = json.loads(student_data.body) if hasattr(student_data, 'body') else student_data
    
    return {
        "report_type": "individual_reports",
        "generated_at": datetime.now().isoformat(),
        "reports": individual_reports,
        "include_charts": include_charts
    }

async def generate_comparison_report(student_ids: List[str], include_charts: bool) -> Dict:
    """
    Generate comparison report for multiple students
    """
    comparison_data = {}
    
    for student_id in student_ids:
        if student_id in student_records:
            student = student_records[student_id]
            history = progress_history.get(student_id, [])
            
            comparison_data[student_id] = {
                "student_info": student,
                "recent_performance": history[-5:] if len(history) >= 5 else history,
                "statistics": {
                    "average_score": student.get('progress', 0),
                    "attendance": student.get('attendance', 100),
                    "trend": calculate_trend(history)
                }
            }
    
    return {
        "report_type": "comparison_report",
        "generated_at": datetime.now().isoformat(),
        "comparison_data": comparison_data,
        "include_charts": include_charts
    }

def create_progress_report_pdf(report_data: Dict, filepath: str):
    """
    Create PDF progress report
    """
    try:
        html_content = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                .header {{ text-align: center; border-bottom: 2px solid #333; padding-bottom: 20px; }}
                .summary {{ background: #f5f5f5; padding: 20px; margin: 20px 0; }}
                .student-table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                .student-table th, .student-table td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
                .student-table th {{ background-color: #f2f2f2; }}
                .status-excellent {{ color: #00ff9d; font-weight: bold; }}
                .status-good {{ color: #64ffda; font-weight: bold; }}
                .status-average {{ color: #ffd166; font-weight: bold; }}
                .status-poor {{ color: #ff6b6b; font-weight: bold; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>Student Progress Report</h1>
                <p>Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
            </div>
            
            <div class="summary">
                <h2>Class Summary</h2>
                <p><strong>Class:</strong> {report_data.get('class_name', 'All Classes')}</p>
                <p><strong>Total Students:</strong> {len(report_data.get('students', {}))}</p>
                <p><strong>Average Progress:</strong> {report_data.get('summary', {}).get('average_progress', 0)}%</p>
            </div>
            
            <h2>Student Progress Details</h2>
            <table class="student-table">
                <thead>
                    <tr>
                        <th>Student Name</th>
                        <th>Class</th>
                        <th>Grade</th>
                        <th>Progress</th>
                        <th>Status</th>
                        <th>Last Updated</th>
                    </tr>
                </thead>
                <tbody>
                    {"".join([f"""
                    <tr>
                        <td>{student['name']}</td>
                        <td>{student['class_name']}</td>
                        <td>{student['grade']}</td>
                        <td>{student['progress']}%</td>
                        <td class="status-{student['status']}">{student['status'].title()}</td>
                        <td>{student['last_updated'][:10]}</td>
                    </tr>
                    """ for student in report_data.get('students', {}).values()])}
                </tbody>
            </table>
        </body>
        </html>
        """
        
        create_pdf_from_html_optional(html_content, filepath)
        
    except Exception as e:
        print(f"Error creating PDF report: {e}")
        raise

def create_progress_report_docx(report_data: Dict, filepath: str):
    """
    Create DOCX progress report
    """
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Title
        doc.add_heading('Student Progress Report', 0)
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        
        # Summary
        doc.add_heading('Class Summary', level=1)
        summary = report_data.get('summary', {})
        doc.add_paragraph(f'Class: {report_data.get("class_name", "All Classes")}')
        doc.add_paragraph(f'Total Students: {len(report_data.get("students", {}))}')
        doc.add_paragraph(f'Average Progress: {summary.get("average_progress", 0)}%')
        doc.add_paragraph(f'Attendance Rate: {summary.get("attendance_rate", 0)}%')
        
        # Student table
        doc.add_heading('Student Progress Details', level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Student Name'
        hdr_cells[1].text = 'Class'
        hdr_cells[2].text = 'Grade'
        hdr_cells[3].text = 'Progress'
        hdr_cells[4].text = 'Status'
        hdr_cells[5].text = 'Last Updated'
        
        # Data rows
        for student in report_data.get('students', {}).values():
            row_cells = table.add_row().cells
            row_cells[0].text = student['name']
            row_cells[1].text = student['class_name']
            row_cells[2].text = student['grade']
            row_cells[3].text = f"{student['progress']}%"
            row_cells[4].text = student['status'].title()
            row_cells[5].text = student['last_updated'][:10]
        
        doc.save(filepath)
        
    except Exception as e:
        print(f"Error creating DOCX report: {e}")
        # Fallback to PDF if DOCX fails
        create_progress_report_pdf(report_data, filepath.replace('.docx', '.pdf'))

# ---------- Initialize Sample Data ----------

def initialize_sample_data():
    """
    Initialize with sample student data for demonstration
    """
    sample_students = [
        {
            "name": "Emma Smith",
            "class_name": "Mathematics 101",
            "grade": "Grade 10",
            "email": "emma.smith@school.edu",
            "progress": 92.5,
            "attendance": 95.0,
            "status": "excellent"
        },
        {
            "name": "Michael Johnson",
            "class_name": "Physics 201",
            "grade": "Grade 11",
            "email": "michael.j@school.edu",
            "progress": 85.0,
            "attendance": 88.0,
            "status": "good"
        },
        {
            "name": "Sarah Davis",
            "class_name": "Chemistry 101",
            "grade": "Grade 10",
            "email": "sarah.d@school.edu",
            "progress": 76.5,
            "attendance": 92.0,
            "status": "average"
        },
        {
            "name": "Robert Wilson",
            "class_name": "Mathematics 101",
            "grade": "Grade 10",
            "email": "robert.w@school.edu",
            "progress": 65.0,
            "attendance": 78.0,
            "status": "poor"
        },
        {
            "name": "Lisa Brown",
            "class_name": "Physics 201",
            "grade": "Grade 11",
            "email": "lisa.b@school.edu",
            "progress": 88.0,
            "attendance": 96.0,
            "status": "good"
        }
    ]
    
    for student_data in sample_students:
        student_id = str(uuid.uuid4())
        student_record = {
            "student_id": student_id,
            **student_data,
            "last_updated": datetime.now().isoformat(),
            "created_at": datetime.now().isoformat()
        }
        student_records[student_id] = student_record
        progress_history[student_id] = []

# ---------- Blackboard Diagrams Generator ----------
@app.post("/generate-visual-aids/")
async def generate_visual_aids(request: Request):
    """
    Generate comprehensive visual aids with actual diagram images
    """
    try:
        data = await request.json()
        topic = data.get("topic", "")
        grade_level = data.get("grade_level", "middle school")
        lesson_duration = data.get("lesson_duration", 45)
        learning_objectives = data.get("learning_objectives", "")
        
        if not topic.strip():
            raise HTTPException(status_code=400, detail="Topic is required")
        
        # First, get the visual aids content
        prompt = f"""
Create a comprehensive set of visual aids for teaching: {topic}

GRADE LEVEL: {grade_level}
LESSON DURATION: {lesson_duration} minutes
LEARNING OBJECTIVES: {learning_objectives}

Please provide a complete visual teaching toolkit including:

1. BLACKBOARD DIAGRAMS (3-5 different types):
   - Main concept diagram
   - Process/flow diagram
   - Comparison chart
   - Timeline or sequence
   - Summary diagram

2. For EACH diagram, provide:
   - Clear title and purpose
   - Step-by-step drawing instructions
   - Simple shapes and elements to use
   - Labeling suggestions
   - Teaching tips

3. VISUAL TEACHING STRATEGIES

Focus on creating practical, classroom-ready visuals.
"""

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": """You are an expert educational consultant specializing in visual teaching methods. 
                    Create comprehensive visual aid packages with clear, implementable diagrams."""
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=3000
        )
        
        visual_aids_content = response.choices[0].message.content
        
        # Generate actual diagram images
        diagram_images = await generate_diagram_images(topic, grade_level)
        
        return JSONResponse(content={
            "topic": topic,
            "grade_level": grade_level,
            "lesson_duration": lesson_duration,
            "visual_aids": visual_aids_content,
            "diagrams": diagram_images,
            "created_at": datetime.now().isoformat()
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Visual aids generation failed: {str(e)}")

async def generate_diagram_images(topic: str, grade_level: str):
    """Generate actual diagram images based on the topic"""
    diagrams = []
    
    try:
        # Generate different types of diagrams
        diagram_types = [
            "main_concept",
            "flowchart", 
            "comparison",
            "timeline",
            "summary"
        ]
        
        for diagram_type in diagram_types:
            image_data = await create_diagram_image(topic, diagram_type, grade_level)
            if image_data:
                diagrams.append({
                    "type": diagram_type,
                    "title": f"{topic} - {diagram_type.replace('_', ' ').title()}",
                    "image_data": image_data,
                    "filename": f"{topic}_{diagram_type}.png"
                })
        
        return diagrams
        
    except Exception as e:
        print(f"Diagram image generation failed: {e}")
        return []

async def create_diagram_image(topic: str, diagram_type: str, grade_level: str):
    """Create a specific diagram image using matplotlib"""
    try:
        # Create a figure with blackboard-like background
        fig, ax = plt.subplots(figsize=(10, 8))
        fig.patch.set_facecolor('#2D5016')  # Blackboard green
        ax.set_facecolor('#2D5016')
        
        # Set up the plot area
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 8)
        ax.axis('off')
        
        # Add title
        title_text = f"{topic}\n{diagram_type.replace('_', ' ').title()}"
        ax.text(5, 7.5, title_text, ha='center', va='center', 
                fontsize=14, color='white', weight='bold',
                bbox=dict(boxstyle="round,pad=0.3", facecolor='#5D4037', alpha=0.8))
        
        # Generate different diagram types
        if diagram_type == "main_concept":
            await draw_main_concept_diagram(ax, topic)
        elif diagram_type == "flowchart":
            await draw_flowchart_diagram(ax, topic)
        elif diagram_type == "comparison":
            await draw_comparison_diagram(ax, topic)
        elif diagram_type == "timeline":
            await draw_timeline_diagram(ax, topic)
        elif diagram_type == "summary":
            await draw_summary_diagram(ax, topic)
        
        # Add watermark
        ax.text(5, 0.3, "EduAssist Visual Aid", ha='center', va='center',
                fontsize=10, color='white', alpha=0.7, style='italic')
        
        # Convert to base64
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight', 
                   facecolor=fig.get_facecolor(), edgecolor='none')
        buffer.seek(0)
        
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        plt.close(fig)
        
        return f"data:image/png;base64,{image_base64}"
        
    except Exception as e:
        print(f"Error creating {diagram_type} diagram: {e}")
        return None

async def draw_main_concept_diagram(ax, topic):
    """Draw a main concept diagram"""
    # Central concept circle
    center_circle = patches.Circle((5, 4), 1.5, fill=True, facecolor='white', 
                                 edgecolor='black', linewidth=2)
    ax.add_patch(center_circle)
    ax.text(5, 4, "Main\nConcept", ha='center', va='center', 
            fontsize=10, weight='bold')
    
    # Surrounding concepts
    concepts = ["Feature 1", "Feature 2", "Feature 3", "Feature 4"]
    angles = [45, 135, 225, 315]
    
    for i, (concept, angle) in enumerate(zip(concepts, angles)):
        rad = np.radians(angle)
        x = 5 + 3 * np.cos(rad)
        y = 4 + 3 * np.sin(rad)
        
        # Connection line
        ax.plot([5 + 1.5 * np.cos(rad), x - 0.5 * np.cos(rad)], 
                [4 + 1.5 * np.sin(rad), y - 0.5 * np.sin(rad)], 
                'w-', linewidth=2)
        
        # Concept box
        box = patches.Rectangle((x-1, y-0.5), 2, 1, fill=True, 
                              facecolor='lightblue', edgecolor='black')
        ax.add_patch(box)
        ax.text(x, y, concept, ha='center', va='center', fontsize=8)

async def draw_flowchart_diagram(ax, topic):
    """Draw a simple flowchart"""
    steps = [
        (5, 6.5, "Start", 'oval'),
        (5, 5.5, "Step 1", 'rectangle'),
        (5, 4.5, "Decision", 'diamond'),
        (3, 3.5, "Option A", 'rectangle'),
        (7, 3.5, "Option B", 'rectangle'),
        (5, 2.5, "End", 'oval')
    ]
    
    for x, y, text, shape in steps:
        if shape == 'oval':
            ellipse = patches.Ellipse((x, y), 2, 0.8, fill=True, 
                                    facecolor='lightgreen', edgecolor='black')
            ax.add_patch(ellipse)
        elif shape == 'rectangle':
            rect = patches.Rectangle((x-1, y-0.4), 2, 0.8, fill=True,
                                  facecolor='lightyellow', edgecolor='black')
            ax.add_patch(rect)
        elif shape == 'diamond':
            diamond = patches.Polygon([(x, y+0.4), (x+1, y), (x, y-0.4), (x-1, y)], 
                                    fill=True, facecolor='lightcoral', edgecolor='black')
            ax.add_patch(diamond)
        
        ax.text(x, y, text, ha='center', va='center', fontsize=8)
    
    # Draw connections
    connections = [
        [(5, 6.1), (5, 5.9)],
        [(5, 5.1), (5, 4.9)],
        [(4, 4.5), (3.5, 4.5), (3.5, 3.9)],
        [(6, 4.5), (6.5, 4.5), (6.5, 3.9)],
        [(3, 3.1), (3, 2.9), (5, 2.9)],
        [(7, 3.1), (7, 2.9), (5, 2.9)]
    ]
    
    for connection in connections:
        x_vals, y_vals = zip(*connection)
        ax.plot(x_vals, y_vals, 'k-', linewidth=2)

async def draw_comparison_diagram(ax, topic):
    """Draw a comparison chart"""
    categories = ["Category A", "Category B", "Category C"]
    values_a = [7, 5, 8]
    values_b = [4, 9, 6]
    
    x = np.arange(len(categories))
    width = 0.35
    
    # Create bars
    bars1 = ax.bar(x - width/2, values_a, width, label='Option 1', 
                  color='skyblue', edgecolor='black')
    bars2 = ax.bar(x + width/2, values_b, width, label='Option 2', 
                  color='lightcoral', edgecolor='black')
    
    # Customize appearance for blackboard
    ax.set_xticks(x)
    ax.set_xticklabels(categories, color='white')
    ax.tick_params(colors='white')
    ax.legend(facecolor='#5D4037', edgecolor='white', labelcolor='white')
    
    # Add value labels on bars
    for bar in bars1:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{height}', ha='center', va='bottom', 
                color='white', weight='bold')
    
    for bar in bars2:
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2., height,
                f'{height}', ha='center', va='bottom', 
                color='white', weight='bold')

async def draw_timeline_diagram(ax, topic):
    """Draw a timeline diagram"""
    events = [
        (1, "Event 1", "Start"),
        (3, "Event 2", "Development"),
        (5, "Event 3", "Progress"),
        (7, "Event 4", "Milestone"),
        (9, "Event 5", "Completion")
    ]
    
    # Draw timeline
    ax.plot([0.5, 9.5], [4, 4], 'w-', linewidth=3)
    
    for x, event, phase in events:
        # Event marker
        ax.plot([x, x], [3.8, 4.2], 'w-', linewidth=2)
        
        # Event circle
        circle = patches.Circle((x, 5), 0.3, fill=True, 
                              facecolor='yellow', edgecolor='black')
        ax.add_patch(circle)
        ax.text(x, 5, event, ha='center', va='center', fontsize=8, weight='bold')
        
        # Phase label
        ax.text(x, 3.5, phase, ha='center', va='top', 
                fontsize=7, color='white', style='italic')

async def draw_summary_diagram(ax, topic):
    """Draw a summary/mind map diagram"""
    # Central topic
    center_circle = patches.Circle((5, 4), 0.8, fill=True, 
                                 facecolor='gold', edgecolor='black')
    ax.add_patch(center_circle)
    ax.text(5, 4, "Summary", ha='center', va='center', 
            fontsize=9, weight='bold')
    
    # Main points
    points = [
        (2, 6, "Key Point 1", 45),
        (8, 6, "Key Point 2", 135),
        (2, 2, "Key Point 3", 315),
        (8, 2, "Key Point 4", 225)
    ]
    
    for x, y, text, angle in points:
        # Connection line
        ax.plot([5, x], [4, y], 'w-', linewidth=2)
        
        # Point box
        box = patches.Rectangle((x-1.2, y-0.4), 2.4, 0.8, fill=True,
                              facecolor='lightgreen', edgecolor='black')
        ax.add_patch(box)
        ax.text(x, y, text, ha='center', va='center', fontsize=8)

# ---------- Enhanced Diagram Generator with Images ----------
@app.post("/generate-diagram/")
async def generate_diagram(request: Request):
    """
    Generate blackboard diagram instructions with actual images
    """
    try:
        data = await request.json()
        topic = data.get("topic", "")
        style = data.get("style", "simple")
        num_diagrams = data.get("num_diagrams", 1)
        content_type = data.get("content_type", "general")
        text_content = data.get("text_content", "")
        
        if not topic.strip():
            raise HTTPException(status_code=400, detail="Topic is required")
        
        # Get diagram instructions
        prompt = f"Create {num_diagrams} simple blackboard diagrams for: {topic}"
        if content_type == "summary" and text_content:
            prompt += f"\nBased on this summary: {text_content}"
        elif content_type == "text_content" and text_content:
            prompt += f"\nBased on this content: {text_content}"
        
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert educator who creates simple, clear blackboard diagrams."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        diagram_instructions = response.choices[0].message.content
        
        # Generate diagram images
        diagram_images = []
        for i in range(min(num_diagrams, 3)):  # Limit to 3 images max
            diagram_type = ["main_concept", "flowchart", "comparison"][i % 3]
            image_data = await create_diagram_image(topic, diagram_type, "middle school")
            if image_data:
                diagram_images.append({
                    "type": diagram_type,
                    "title": f"{topic} - Diagram {i+1}",
                    "image_data": image_data,
                    "filename": f"{topic}_diagram_{i+1}.png"
                })
        
        return JSONResponse(content={
            "topic": topic,
            "style": style,
            "num_diagrams": num_diagrams,
            "content_type": content_type,
            "instructions": diagram_instructions,
            "diagrams": diagram_images,
            "created_at": datetime.now().isoformat()
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Diagram generation failed: {str(e)}")

# ---------- Download Diagram Images ----------
@app.post("/download-diagrams/")
async def download_diagrams(request: Request):
    """
    Download generated diagrams as a ZIP file
    """
    try:
        data = await request.json()
        topic = data.get("topic", "")
        diagrams = data.get("diagrams", [])
        
        if not diagrams:
            raise HTTPException(status_code=400, detail="No diagrams to download")
        
        # Create a unique directory for this download
        job_id = str(uuid.uuid4())
        output_dir = os.path.join(OUTPUT_FOLDER, f"diagrams_{job_id}")
        os.makedirs(output_dir, exist_ok=True)
        
        # Save each diagram
        saved_files = []
        for diagram in diagrams:
            if diagram.get('image_data'):
                # Extract base64 data
                image_data = diagram['image_data'].split(',')[1]
                image_bytes = base64.b64decode(image_data)
                
                filename = diagram.get('filename', f"{topic}_diagram.png")
                file_path = os.path.join(output_dir, filename)
                
                with open(file_path, 'wb') as f:
                    f.write(image_bytes)
                saved_files.append(file_path)
        
        # Create ZIP file
        zip_filename = f"{topic}_diagrams.zip"
        zip_path = os.path.join(OUTPUT_FOLDER, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for file_path in saved_files:
                zipf.write(file_path, os.path.basename(file_path))
        
        # Clean up individual files
        for file_path in saved_files:
            try:
                os.remove(file_path)
            except:
                pass
        try:
            os.rmdir(output_dir)
        except:
            pass
        
        return JSONResponse(content={
            "download_url": f"/download/{zip_filename}",
            "filename": zip_filename,
            "diagram_count": len(diagrams),
            "message": f"Download ready for {len(diagrams)} diagrams"
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Download creation failed: {str(e)}")

# ---------- Simple Diagram Creator ----------
@app.post("/create-simple-diagram/")
async def create_simple_diagram(request: Request):
    """
    Create a simple diagram with custom elements
    """
    try:
        data = await request.json()
        topic = data.get("topic", "")
        diagram_type = data.get("diagram_type", "concept_map")
        elements = data.get("elements", [])
        
        # Generate diagram based on type
        image_data = await create_custom_diagram(topic, diagram_type, elements)
        
        return JSONResponse(content={
            "topic": topic,
            "diagram_type": diagram_type,
            "image_data": image_data,
            "filename": f"{topic}_{diagram_type}.png",
            "created_at": datetime.now().isoformat()
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Simple diagram creation failed: {str(e)}")

async def create_custom_diagram(topic: str, diagram_type: str, elements: list):
    """Create a custom diagram based on user elements"""
    try:
        fig, ax = plt.subplots(figsize=(10, 8))
        fig.patch.set_facecolor('#2D5016')
        ax.set_facecolor('#2D5016')
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 8)
        ax.axis('off')
        
        # Add title
        ax.text(5, 7.5, f"{topic}\nCustom Diagram", ha='center', va='center', 
                fontsize=14, color='white', weight='bold',
                bbox=dict(boxstyle="round,pad=0.3", facecolor='#5D4037', alpha=0.8))
        
        # Add sample content based on diagram type
        if diagram_type == "concept_map":
            await draw_main_concept_diagram(ax, topic)
        elif diagram_type == "flowchart":
            await draw_flowchart_diagram(ax, topic)
        elif diagram_type == "timeline":
            await draw_timeline_diagram(ax, topic)
        
        # Convert to base64
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=150, bbox_inches='tight',
                   facecolor=fig.get_facecolor(), edgecolor='none')
        buffer.seek(0)
        
        image_base64 = base64.b64encode(buffer.getvalue()).decode()
        plt.close(fig)
        
        return f"data:image/png;base64,{image_base64}"
        
    except Exception as e:
        print(f"Error creating custom diagram: {e}")
        return None

# ---------- Knowledge Base Endpoints ----------
@app.post("/explain-concept/")
async def explain_concept(request: Request):
    """
    Explain complex concepts with simple explanations, analogies, and local language support
    """
    try:
        data = await request.json()
        concept = data.get("concept", "")
        grade_level = data.get("grade_level", "general")
        language = data.get("language", "english")
        include_analogies = data.get("include_analogies", True)
        include_examples = data.get("include_examples", True)
        include_visuals = data.get("include_visuals", False)
        
        if not concept.strip():
            raise HTTPException(status_code=400, detail="Concept is required")
        
        # Create explanation prompt
        prompt = f"""
Explain the concept: "{concept}"

GRADE LEVEL: {grade_level}
LANGUAGE: {language}
INCLUDE ANALOGIES: {include_analogies}
INCLUDE EXAMPLES: {include_examples}
INCLUDE VISUALS: {include_visuals}

Please provide a comprehensive explanation with:

1. SIMPLE DEFINITION: Easy-to-understand definition
2. KEY POINTS: 3-5 main characteristics or components
3. {f"ANALOGIES: 2-3 relatable analogies from everyday life" if include_analogies else ""}
4. {f"EXAMPLES: Practical examples and applications" if include_examples else ""}
5. COMMON MISCONCEPTIONS: What people often get wrong
6. REAL-WORLD CONNECTIONS: How this applies to daily life
7. {f"VISUAL DESCRIPTION: Simple way to visualize this concept" if include_visuals else ""}
8. RELATED CONCEPTS: What to learn next

Make it:
- Age-appropriate for {grade_level}
- Clear and engaging
- Practical and relatable
- Culturally appropriate for {language} speakers
- Free of unnecessary jargon
"""

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": f"""You are an expert educator who explains complex concepts in simple, engaging ways.
                    You specialize in making difficult topics accessible to all age groups.
                    You provide relatable analogies and practical examples.
                    You adapt your explanations to be culturally relevant."""
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=3000,
            top_p=0.9
        )
        
        explanation_content = response.choices[0].message.content
        
        # Generate related concepts
        related_concepts = await generate_related_concepts(concept, grade_level)
        
        return JSONResponse(content={
            "status": "success",
            "concept": concept,
            "grade_level": grade_level,
            "language": language,
            "explanation": explanation_content,
            "related_concepts": related_concepts,
            "created_at": datetime.now().isoformat()
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Concept explanation failed: {str(e)}")

async def generate_related_concepts(concept: str, grade_level: str):
    """Generate related concepts for further learning"""
    try:
        prompt = f"""
Given the concept "{concept}" for {grade_level} level, suggest 5-7 related concepts that would help build deeper understanding.
Return only a comma-separated list without additional text.

Focus on:
- Prerequisite concepts
- Complementary topics
- Advanced extensions
- Practical applications
"""

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": "Generate relevant, educational concepts that build on the main topic. Return only a comma-separated list."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=200
        )
        
        concepts_text = response.choices[0].message.content.strip()
        concepts_list = [c.strip() for c in concepts_text.split(',') if c.strip()]
        
        return concepts_list[:7]  # Limit to 7 concepts
        
    except Exception as e:
        print(f"Related concepts generation failed: {e}")
        return []

@app.post("/compare-concepts/")
async def compare_concepts(request: Request):
    """
    Compare two or more concepts with similarities and differences
    """
    try:
        data = await request.json()
        concepts = data.get("concepts", [])
        comparison_focus = data.get("focus", "similarities and differences")
        grade_level = data.get("grade_level", "general")
        
        if len(concepts) < 2:
            raise HTTPException(status_code=400, detail="At least two concepts are required for comparison")
        
        concepts_text = ", ".join(concepts)
        
        prompt = f"""
Compare these concepts: {concepts_text}

COMPARISON FOCUS: {comparison_focus}
GRADE LEVEL: {grade_level}

Please provide a comprehensive comparison with:

1. BASIC DEFINITIONS: Simple definition of each concept
2. KEY SIMILARITIES: What they have in common
3. KEY DIFFERENCES: How they are distinct
4. RELATIONSHIPS: How they connect or interact
5. WHEN TO USE EACH: Practical applications and contexts
6. COMMON CONFUSIONS: What people mix up
7. LEARNING TIPS: How to remember the differences

Make the comparison:
- Clear and structured
- Practical and relevant
- Age-appropriate
- Helpful for understanding
"""

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert at comparing and contrasting concepts. You highlight both similarities and differences in clear, educational ways."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,
            max_tokens=2500
        )
        
        comparison_content = response.choices[0].message.content
        
        return JSONResponse(content={
            "status": "success",
            "concepts": concepts,
            "comparison_focus": comparison_focus,
            "grade_level": grade_level,
            "comparison": comparison_content,
            "created_at": datetime.now().isoformat()
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Concept comparison failed: {str(e)}")

@app.post("/generate-analogies/")
async def generate_analogies(request: Request):
    """
    Generate multiple analogies for a concept
    """
    try:
        data = await request.json()
        concept = data.get("concept", "")
        context = data.get("context", "general")
        num_analogies = data.get("num_analogies", 3)
        grade_level = data.get("grade_level", "general")
        
        if not concept.strip():
            raise HTTPException(status_code=400, detail="Concept is required")
        
        prompt = f"""
Generate {num_analogies} creative and relatable analogies for: "{concept}"

CONTEXT: {context}
GRADE LEVEL: {grade_level}

For each analogy, provide:
1. The analogy itself
2. How it relates to the concept
3. Why it helps understanding
4. Potential limitations

Focus on analogies that are:
- Easy to understand
- Culturally relevant
- Age-appropriate
- Memorable and engaging
"""

        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": "You are a master of creating perfect analogies that make complex concepts instantly understandable."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.8,
            max_tokens=2000
        )
        
        analogies_content = response.choices[0].message.content
        
        return JSONResponse(content={
            "status": "success",
            "concept": concept,
            "context": context,
            "grade_level": grade_level,
            "analogies": analogies_content,
            "created_at": datetime.now().isoformat()
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Analogy generation failed: {str(e)}")

# ---------- Download Single File ----------
@app.get("/download/{filename}")
async def download(filename: str):
    file_path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found.")
    
    # Determine media type based on file extension
    if filename.lower().endswith('.pdf'):
        media_type = "application/pdf"
    elif filename.lower().endswith('.docx'):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        media_type = "application/octet-stream"
    
    return FileResponse(
        file_path, 
        media_type=media_type, 
        filename=filename,
        headers={'Content-Disposition': f'attachment; filename="{filename}"'}
    )

# ---------- Conversations ----------
@app.get("/conversation/{conversation_id}")
async def get_conversation(conversation_id: str):
    if conversation_id not in conversations:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return JSONResponse(content=conversations[conversation_id])

@app.get("/conversations/")
async def list_conversations():
    return JSONResponse(content=[
        {
            "id": cid,
            "created_at": conv["created_at"],
            "message_count": len(conv["messages"]),
            "lesson_plan_count": len(conv["lesson_plans"])
        }
        for cid, conv in conversations.items()
    ])

@app.delete("/conversation/{conversation_id}")
async def delete_conversation(conversation_id: str):
    if conversation_id in conversations:
        del conversations[conversation_id]
        return {"message": "Conversation deleted successfully"}
    raise HTTPException(status_code=404, detail="Conversation not found")

# ---------- Health ----------
@app.get("/health")
async def health_check():
    return {"status": "healthy", "timestamp": datetime.now().isoformat()}

# ---------- Cleanup Old Files ----------
def cleanup_old_files():
    """Clean up files older than 24 hours"""
    try:
        current_time = time.time()
        for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER]:
            if not os.path.exists(folder):
                continue
            for filename in os.listdir(folder):
                file_path = os.path.join(folder, filename)
                if os.path.isfile(file_path):
                    # Delete files older than 24 hours
                    if current_time - os.path.getctime(file_path) > 24 * 3600:
                        os.remove(file_path)
                        print(f"Cleaned up old file: {filename}")
    except Exception as e:
        print(f"Cleanup error: {e}")

def run_scheduler():
    while True:
        schedule.run_pending()
        time.sleep(3600)  # Check every hour

# Schedule cleanup to run daily
schedule.every().day.at("03:00").do(cleanup_old_files)

# Start cleanup thread
cleanup_thread = threading.Thread(target=run_scheduler, daemon=True)
cleanup_thread.start()

# ---------- Errors ----------
@app.exception_handler(404)
async def not_found_handler(request, exc):
    return JSONResponse(status_code=404, content={"error": "Endpoint not found", "detail": str(exc)})

@app.exception_handler(500)
async def internal_error_handler(request, exc):
    return JSONResponse(status_code=500, content={"error": "Internal server error", "detail": str(exc)})

# ---------- Startup ----------
@app.on_event("startup")
async def startup_event():
    print("🚀 EduAssist (Groq) Backend started successfully!")
    # Initial cleanup
    cleanup_old_files()
    # Initialize sample data
    initialize_sample_data()

# if __name__ == "__main__":
#     import uvicorn
#     uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)

if __name__ == "__main__":
    app.run()
